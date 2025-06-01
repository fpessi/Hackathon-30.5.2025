from docx import Document
from docx.shared import Pt, RGBColor
import os
import sys

def get_output_directory():
    """Ensures the output directory exists and returns its path

    Returns:
        str: filepath
    """
    if '__file__' in globals():
        base_path = os.path.dirname(__file__)
    else:
        base_path = os.path.dirname(os.path.abspath(sys.argv[0]))

    relative_path = os.path.join(base_path, "..", "Case")
    os.makedirs(relative_path, exist_ok=True)
    return relative_path

def style_run(run, size=14, bold=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def make_report(general_data, field_work_data, supply_data, special_info):
    """Makes a docx file based on input

    Args:
        general_data (dir): data for document
        field_work_data (dir): data for document
        supply_data (dir): data for document
        special_info (dir): data for document
    """
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run('Technical Field Report')
    style_run(run, size=28, bold=True)

    # Section: General Information
    general_header = doc.add_paragraph()
    run = general_header.add_run('General Information')
    style_run(run, size=20, bold=True)

    table1 = doc.add_table(rows=len(general_data), cols=2)
    table1.style = 'Table Grid'
    for i, (label, value) in enumerate(general_data.items()):
        cell_left = table1.cell(i, 0)
        cell_left.text = label
        style_run(cell_left.paragraphs[0].runs[0])
        cell_right = table1.cell(i, 1)
        cell_right.text = value
        style_run(cell_right.paragraphs[0].runs[0])

    # Section: Field Work Description
    field_header = doc.add_paragraph()
    run = field_header.add_run('Field Work Description')
    style_run(run, size=20, bold=True)

    table2 = doc.add_table(rows=len(field_work_data), cols=2)
    table2.style = 'Table Grid'
    for i, (label, value) in enumerate(field_work_data.items()):
        cell_left = table2.cell(i, 0)
        cell_left.text = label
        style_run(cell_left.paragraphs[0].runs[0])
        cell_right = table2.cell(i, 1)
        cell_right.text = value
        style_run(cell_right.paragraphs[0].runs[0])

    # Section: Detailed Supply Information
    supply_header = doc.add_paragraph()
    run = supply_header.add_run('Detailed Supply Information')
    style_run(run, size=20, bold=True)

    table3 = doc.add_table(rows=len(supply_data), cols=2)
    table3.style = 'Table Grid'
    for i, (label, value) in enumerate(supply_data.items()):
        cell_left = table3.cell(i, 0)
        cell_left.text = label
        style_run(cell_left.paragraphs[0].runs[0])
        cell_right = table3.cell(i, 1)
        cell_right.text = value
        style_run(cell_right.paragraphs[0].runs[0])

    # Section: Special Information
    special_header = doc.add_paragraph()
    run = special_header.add_run('Special Information')
    style_run(run, size=20, bold=True)

    table4 = doc.add_table(rows=1, cols=1)
    table4.style = 'Table Grid'
    cell = table4.cell(0, 0)
    cell.text = special_info
    style_run(cell.paragraphs[0].runs[0])

    # Save the document
    output_dir = get_output_directory()
    output_path = os.path.join(output_dir, "field_report.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

# Example usage
"""
if __name__ == "__main__":
    general_data = {
        "Location": "Espoo headquarters",
        "Date": "2025-06-01",
        "Service personnel": "John Johnson",
        "Equipment": "Excavator ZX200",
        "Time of service": "08:00 - 12:30"
    }

    field_work_data = {
        "Reason of service": "Scheduled maintenance",
        "Work description": "Replaced hydraulic filters and checked fluid levels.",
        "Problems with service": "None",
        "Notes": "Unit is in good condition."
    }

    supply_data = {
        "Supplies and amounts": "2x Hydraulic filter, 1L Oil",
        "Weight": "3 kg",
        "Length": "30 cm",
        "Width": "10 cm"
    }

    special_info = "Next maintenance due in 3 months."

    make_report(general_data, field_work_data, supply_data, special_info)
"""